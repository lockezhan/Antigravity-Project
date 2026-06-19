import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_run(p, text, font_name='SimSun', size_pt=12, bold=False):
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    return run

def create_paper():
    doc = Document()
    
    # 页面设置 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # 页眉设置
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(header_p, "第二十一届中国研究生电子设计竞赛\n", "SimSun", 10.5)
    add_run(header_p, "面向异构计算平台的高性能物理信息神经网络（PINN）流体力学求解器软硬件协同设计", "SimSun", 10.5)
    
    # 封面
    doc.add_paragraph() # spacing
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "第二十一届中国研究生电子设计竞赛\n\n", "SimHei", 26) # 一号 26pt
    add_run(p, "技术论文\n\n\n", "SimHei", 26)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "论文题目：\n", "SimHei", 18) # 小二号 18pt
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_title, "（中文）面向异构计算平台的高性能物理信息神经网络（PINN）流体力学求解器软硬件协同设计\n", "SimHei", 18)
    add_run(p_title, "（英文）Hardware-Software Co-Design of High-Performance PINN Solver\n\n\n", "SimHei", 18)
    
    p = doc.add_paragraph()
    add_run(p, "参赛单位：[填入参赛单位]\n", "SimHei", 15) # 小三号 15pt
    add_run(p, "队伍名称：[填入队伍名称]\n", "SimHei", 15)
    add_run(p, "指导老师：[填入指导老师]\n", "SimHei", 15)
    add_run(p, "参赛队员：[填入参赛队员]\n", "SimHei", 15)
    add_run(p, "完成时间：[填入完成时间]\n", "SimHei", 15)
    
    doc.add_page_break()
    
    # 中英文摘要 (正文格式: 宋体小四 12pt, 行距 20pt)
    def add_body_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, bold=False):
        p = doc.add_paragraph()
        p.alignment = align
        if first_line_indent:
            p.paragraph_format.first_line_indent = Pt(24) # 2 chars
        p.paragraph_format.line_spacing = Pt(20) # 固定行间距 20pt
        p.paragraph_format.line_spacing_rule = 4 # WD_LINE_SPACING.EXACTLY
        parts = text.split("**")
        for i, part in enumerate(parts):
            if part:
                is_bold = bold or (i % 2 != 0)
                add_run(p, part, "SimSun", 12, bold=is_bold)
        return p

    def add_heading(text, level):
        p = doc.add_heading(level=level)
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = add_run(p, text, "SimHei", 18) # 小二号 18pt
        elif level == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = add_run(p, text, "SimHei", 14) # 小三号 14pt (approx)
        elif level == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = add_run(p, text, "SimHei", 12) # 小四号 12pt
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.line_spacing_rule = 4
        return p

    with open('/mnt/c/Users/elite/.gemini/antigravity/brain/56aafa08-3008-4ac5-8d46-71e810b295fe/研电赛技术论文.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    idx = 0
    # Read abstracts
    while idx < len(lines):
        line = lines[idx].strip()
        if line == '## 目录':
            break
        if line:
            if line.startswith('**中文摘要：**') or line.startswith('**英文摘要：**') or line.startswith('**关键字：**'):
                add_body_paragraph(line, first_line_indent=False)
            else:
                add_body_paragraph(line)
        idx += 1
        
    doc.add_page_break()

    # TOC
    p_toc = doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_toc, "目 录", "SimHei", 16) # 三号
    
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    doc.add_page_break()
    
    # 页脚设置 (页码从正文开始)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    # Process Chapters
    while idx < len(lines):
        line = lines[idx].strip()
        if line.startswith('## 第') or line.startswith('## 参考文献'):
            title = line.replace('## ', '')
            if not line.startswith('## 第1章'):
                doc.add_page_break()
            add_heading(title, 1)
        elif line.startswith('## 数据与代码'):
            title = line.replace('## ', '')
            doc.add_page_break()
            add_heading(title, 1)
        elif line.startswith('### '):
            title = line.replace('### ', '')
            add_heading(title, 2)
        elif line.startswith('#### '):
            title = line.replace('#### ', '')
            add_heading(title, 3)
        elif line.startswith('- ') or line.startswith('* '):
            add_body_paragraph(line, first_line_indent=False)
        elif line.startswith('![') and line.endswith(')'):
            # image insertion
            # ![alt](file:///path)
            try:
                img_path = line.split('(')[1].replace(')', '').replace('file:///', '')
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_img = p_img.add_run()
                    r_img.add_picture(img_path, width=Cm(14))
            except Exception as e:
                pass
        elif line:
            if not line.startswith('---') and not line.startswith('```') and not line.startswith('## 目录'):
                add_body_paragraph(line)
        idx += 1
        
    output_dir = '/home/elite/workspace'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    doc.save(os.path.join(output_dir, 'paper_output.docx'))
    print("Done generating docx.")

if __name__ == '__main__':
    create_paper()
